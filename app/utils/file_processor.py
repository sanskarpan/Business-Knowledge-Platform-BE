import os
import hashlib
import mimetypes
from typing import Tuple, Optional, List
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import markdown
from PIL import Image
from app.config import settings

def validate_file(file_content: bytes, filename: str) -> Tuple[bool, Optional[str]]:
    """Validate file type and size"""
    # Check file size
    if len(file_content) > settings.max_file_size:
        return False, f"File size exceeds {settings.max_file_size / (1024*1024):.0f}MB limit"
    
    # Check file type by extension and content
    allowed_types = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
    }
    
    file_ext = os.path.splitext(filename)[1].lower()
    
    if file_ext not in allowed_types:
        return False, f"File type {file_ext} not supported. Allowed: {', '.join(allowed_types.keys())}"
    
    # Basic MIME type validation (simplified since python-magic might not be available)
    try:
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type and not mime_type.startswith(('application/', 'text/', 'image/')):
            return False, f"Invalid file type: {mime_type}"
    except Exception:
        pass  # Skip MIME validation if it fails
    
    return True, None

def extract_text_content(file_path: str, file_type: str) -> str:
    """Extract text content from various file types"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return extract_pdf_text(file_path)
        elif file_ext == '.docx':
            return extract_docx_text(file_path)
        elif file_ext in ['.txt', '.md']:
            return extract_text_file(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            return extract_image_text(file_path)
        else:
            return ""
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return f"Error extracting content: {str(e)}"

def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text + "\n"
                except Exception as e:
                    print(f"Error extracting page {page_num + 1}: {e}")
                    continue
        return text.strip()
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        text = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text.strip()
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_text_file(file_path: str) -> str:
    """Extract text from TXT or MD file"""
    try:
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    
                # If it's markdown, convert to plain text
                if file_path.lower().endswith('.md'):
                    # Simple markdown to text conversion
                    import re
                    # Remove markdown syntax
                    content = re.sub(r'#{1,6}\s+', '', content)  # Headers
                    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Bold
                    content = re.sub(r'\*(.*?)\*', r'\1', content)  # Italic
                    content = re.sub(r'`(.*?)`', r'\1', content)  # Code
                    content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)  # Links
                
                return content.strip()
            except UnicodeDecodeError:
                continue
        
        return "Error: Could not decode file with any supported encoding"
    except Exception as e:
        return f"Error reading text file: {str(e)}"

def extract_image_text(file_path: str) -> str:
    """Extract text from image using OCR (requires pytesseract and tesseract)"""
    try:
        # Check if tesseract is available
        try:
            import pytesseract
        except ImportError:
            return "OCR not available: pytesseract not installed"
        
        # Open and process image
        with Image.open(file_path) as image:
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using OCR
            extracted_text = pytesseract.image_to_string(image)
            
            if extracted_text.strip():
                return extracted_text.strip()
            else:
                return "No text found in image"
                
    except ImportError:
        return "OCR not available: tesseract not installed"
    except Exception as e:
        return f"Error extracting text from image: {str(e)}"

def generate_file_hash(file_content: bytes) -> str:
    """Generate SHA-256 hash of file content"""
    return hashlib.sha256(file_content).hexdigest()

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into chunks for vector embedding"""
    if not text or not text.strip():
        return []
    
    chunks = []
    text = text.strip()
    text_length = len(text)
    
    if text_length <= chunk_size:
        return [text]
    
    start = 0
    
    while start < text_length:
        end = start + chunk_size
        
        # If this is not the last chunk, try to break at a sentence or word boundary
        if end < text_length:
            last_period = text.rfind('.', start, end)
            last_newline = text.rfind('\n', start, end)
            last_space = text.rfind(' ', start, end)
            
            # Choose the best break point
            break_point = max(last_period, last_newline)
            if break_point > start + chunk_size - 200:  # If break point is reasonably close to end
                end = break_point + 1
            elif last_space > start + chunk_size - 100:  # Fall back to word boundary
                end = last_space
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position considering overlap
        if end >= text_length:
            break
        
        # For overlap, move back from the end position
        start = max(start + 1, end - overlap)
    
    return chunks

def get_file_metadata(file_path: str, file_content: bytes) -> dict:
    """Extract basic file metadata"""
    try:
        stat = os.stat(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        metadata = {
            'file_size': len(file_content),
            'file_extension': file_ext,
            'file_hash': generate_file_hash(file_content),
            'created_at': stat.st_ctime,
            'modified_at': stat.st_mtime,
        }
        
        # Add file-type specific metadata
        if file_ext == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    metadata['page_count'] = len(pdf_reader.pages)
                    if pdf_reader.metadata:
                        metadata['pdf_title'] = pdf_reader.metadata.get('/Title', '')
                        metadata['pdf_author'] = pdf_reader.metadata.get('/Author', '')
            except Exception:
                pass
        
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            try:
                with Image.open(file_path) as image:
                    metadata['image_width'] = image.width
                    metadata['image_height'] = image.height
                    metadata['image_format'] = image.format
            except Exception:
                pass
        
        return metadata
    except Exception as e:
        return {'error': f"Could not extract metadata: {str(e)}"}

def cleanup_temp_file(file_path: str) -> bool:
    """Clean up temporary files"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            return True
        return False
    except Exception as e:
        print(f"Error cleaning up file {file_path}: {e}")
        return False

# Utility function for testing file processing
def process_sample_content():
    """Process the sample content provided in the project brief"""
    
    # Sample bill content (from the images)
    sample_bill_content = """
    BEAN & BREW COFFEE SHOP
    123 Main Street, San Francisco
    Tel: (415) 555-0123
    
    RECEIPT #: 6721
    DATE: 30/04/2025
    TIME: 08:37 AM
    SERVER: Miguel
    TABLE: 4
    GUESTS: 2
    
    1 x Cappuccino         $5.50
    1 x Americano          $4.25
    1 x Avocado Toast      $9.95
    1 x Granola Bowl       $8.75
    
    SUBTOTAL:              $28.45
    TAX (8.5%):            $2.42
    SERVICE CHARGE (18%):   $5.12
    
    TOTAL:                 $35.99
    
    PAYMENT METHOD: Credit Card
    CARD TYPE: VISA ****1234
    
    TIP: $7.00
    TOTAL PAID: $42.99
    
    Thank you for visiting!
    COME AGAIN SOON
    """
    
    # Sample feedback content
    sample_feedback_content = """
    Customer Feedback - Bean & Brew Coffee Shop
    
    Feedback #1 - April 5, 2025
    Name: Emma Rodriguez
    Rating: ★★★★★ (5/5)
    Visit Date: April 5, 2025
    
    I absolutely love the atmosphere at Bean & Brew! The avocado toast is to die for, 
    and the cappuccino was perfectly made with beautiful latte art. The staff was friendly 
    and attentive, especially Miguel who remembered my name from my last visit. This place 
    has become my go-to weekend spot. Highly recommend the window seats for people watching!
    """
    
    return {
        'bill_chunks': chunk_text(sample_bill_content),
        'feedback_chunks': chunk_text(sample_feedback_content),
        'bill_metadata': {
            'type': 'receipt',
            'business': 'Bean & Brew Coffee Shop',
            'date': '2025-04-30',
            'total': '$42.99'
        }
    }